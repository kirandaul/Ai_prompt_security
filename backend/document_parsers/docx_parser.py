"""
DOCX (Word) document parser using python-docx library.
Extracts text from paragraphs, tables, etc.
"""

from typing import Dict, List, Any
from .base_parser import BaseDocumentParser

try:
    from docx import Document
except ImportError:
    Document = None


class DOCXParser(BaseDocumentParser):
    """Parse DOCX (Word) documents."""
    
    async def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX and extract text with section information.
        
        Returns:
            {
                "success": bool,
                "filename": str,
                "file_type": "docx",
                "pages": [
                    {
                        "section": "paragraph" or "table",
                        "section_num": int,
                        "text": "...",
                        "length": int
                    },
                    ...
                ],
                "total_pages": int,  # Number of sections
                "text": "combined text",
                "metadata": {...}
            }
        """
        
        if not Document:
            return {
                "success": False,
                "error": "python-docx library not installed",
                "filename": filename,
                "file_type": "docx"
            }
        
        try:
            import io
            doc = Document(io.BytesIO(file_content))
            
            pages = []
            all_text = []
            section_num = 0
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    section_num += 1
                    pages.append({
                        "section": "paragraph",
                        "section_num": section_num,
                        "text": para.text,
                        "length": len(para.text)
                    })
                    all_text.append(f"--- Section {section_num} (Paragraph) ---\n{para.text}")
            
            # Extract from tables
            for table_idx, table in enumerate(doc.tables, start=1):
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        section_num += 1
                        pages.append({
                            "section": "table",
                            "section_num": section_num,
                            "table_idx": table_idx,
                            "row_idx": row_idx,
                            "text": row_text,
                            "length": len(row_text)
                        })
                        all_text.append(f"--- Section {section_num} (Table {table_idx}, Row {row_idx}) ---\n{row_text}")
            
            # Extract metadata
            metadata = {}
            try:
                core_props = doc.core_properties
                metadata = {
                    "title": core_props.title or "",
                    "subject": core_props.subject or "",
                    "author": core_props.author or "",
                    "comments": core_props.comments or "",
                    "category": core_props.category or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                }
            except:
                pass
            
            return {
                "success": True,
                "filename": filename,
                "file_type": "docx",
                "pages": pages,
                "total_pages": len(pages),
                "text": "\n".join(all_text),
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "file_type": "docx",
                "error": str(e),
                "pages": [],
                "total_pages": 0,
                "text": ""
            }
